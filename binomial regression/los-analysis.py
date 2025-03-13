import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.genmod.families.family import NegativeBinomial
import tkinter as tk
from tkinter import filedialog
from scipy import stats
import docx
from docx import Document
from docx.shared import Inches
from io import BytesIO
import statsmodels.discrete.discrete_model as discrete
from statsmodels.regression.mixed_linear_model import MixedLM

def select_file(title, file_types, save=False):
    """Allow user to select a file"""
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    
    try:
        if save:
            file_path = filedialog.asksaveasfilename(
                title=title,
                filetypes=file_types,
                defaultextension=file_types[0][1]
            )
        else:
            file_path = filedialog.askopenfilename(
                title=title,
                filetypes=file_types
            )
    finally:
        root.destroy()
    
    return file_path if file_path else None

# Allow user to select input file
print("Please select the input Excel file...")
file_path = select_file(
    "Select Excel Data File", 
    [("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
)

if not file_path:
    print("No file selected. Exiting.")
    exit()

# Import the data
print(f"Loading data from: {file_path}")
df = pd.read_excel(file_path, sheet_name="Stata")

# Setup
pd.set_option('display.max_columns', None)

# Encode categorical variables if not already encoded
categorical_vars = ['sex', 'marital_status', 'employment_status', 'purpose', 'accomd_type', 'us_state']
encoded_vars = {}

for var in categorical_vars:
    if var in df.columns:
        # Check if variable is already numeric
        if not pd.api.types.is_numeric_dtype(df[var]):
            new_var = f"{var}_enc"
            df[new_var] = pd.Categorical(df[var]).codes
            encoded_vars[var] = new_var
        else:
            encoded_vars[var] = var

# Set the truncation point for los (assuming truncation at 0)
df['los_trunc'] = df['los'].copy()
df.loc[df['los_trunc'] <= 0, 'los_trunc'] = np.nan

# Check for missing data
print("\nMissing data summary:")
missing_data_summary = df.isnull().sum()
print(missing_data_summary)

print("\nMissing data patterns:")
missing_patterns = df.isnull().sum(axis=1)
missing_patterns_counts = missing_patterns.value_counts().sort_index()
print(missing_patterns_counts)

# Visualize los distribution
plt.figure(figsize=(10, 6))
sns.histplot(df['los_trunc'], discrete=True)
plt.title('Histogram of Length of Stay')
plt.tight_layout()
los_hist_img = BytesIO()
plt.savefig(los_hist_img, format='png')
los_hist_img.seek(0)
plt.close()

# Summarize los by purpose
purpose_stats = None
if 'purpose_enc' in df.columns:
    print("\nLength of stay by purpose:")
    purpose_stats = df.groupby('purpose_enc')['los_trunc'].agg(['count', 'mean', 'median', 'min', 'max', 'std'])
    print(purpose_stats)

# Detailed summary of los_trunc
print("\nDetailed summary of length of stay:")
los_describe = df['los_trunc'].describe(percentiles=[.25, .5, .75, .90, .95, .99])
print(los_describe)

# Cleaning process
# Step 1: Drop missing datapoints for key variables
key_vars = ['los', 'immigrant_population', 'import_from_slu', 'age', 
            encoded_vars.get('sex', 'sex_enc'), 
            encoded_vars.get('marital_status', 'marital_status_enc'), 
            encoded_vars.get('employment_status', 'employment_status_enc'), 
            'distance_miles', 
            encoded_vars.get('purpose', 'purpose_enc'), 
            encoded_vars.get('accomd_type', 'accomd_type_enc'), 
            'month_travel', 'state_percapita_income', 'state_unemployment']

# Count missing values per row for key variables
df['missing'] = df[key_vars].isnull().sum(axis=1)
print("\nNumber of missing values per observation:")
missing_values_count = df['missing'].value_counts().sort_index()
print(missing_values_count)

# Drop observations with missing values in key variables
df_clean = df[df['missing'] == 0].drop('missing', axis=1)
print(f"\nRemaining observations after dropping missing values: {len(df_clean)}")

# Step 2: Drop outliers in length of stay
los_p95 = np.percentile(df_clean['los_trunc'].dropna(), 95)
df_clean['los_capped'] = df_clean['los_trunc'].copy()
df_clean.loc[df_clean['los_capped'] > los_p95, 'los_capped'] = los_p95

# Visualize the capped los distribution
plt.figure(figsize=(10, 6))
sns.histplot(df_clean['los_capped'], discrete=True)
plt.title('Histogram of Capped Length of Stay')
plt.tight_layout()
los_capped_img = BytesIO()
plt.savefig(los_capped_img, format='png')
los_capped_img.seek(0)
plt.close()

# Step 3: Clean up the purpose of trip column
# Create a new simplified purpose variable
purpose_mapping = {
    1: 1,  # BUSINESS/MEETING -> Business
    2: 1,  # CONVENTION -> Business
    3: 1,  # CREW -> Business
    5: 2,  # EVENT -> Events
    6: 2,  # EVENTS -> Events
    7: 4,  # HONEYMOON -> Pleasure
    8: 5,  # INTRANSIT PASSEN -> Other
    9: 5,  # OTHER -> Other
    10: 4, # PLEASURE/HOLIDAY -> Pleasure
    11: 5, # RESIDENT -> Other
    12: 2, # SAINT LUCIA CARN -> Events
    13: 2, # SAINT LUCIA JAZZ -> Events
    14: 5, # SPORTS -> Other
    15: 5, # STUDY -> Other
    16: 5, # VISITING FRIENDS -> Other
    17: 3, # WEDDING -> Wedding
    18: 4, # pLEASURE/HOLIDAY -> Pleasure
    4: 5,  # CRICKET -> Other
}

purpose_labels = {
    1: "Business",
    2: "Events",
    3: "Wedding",
    4: "Pleasure",
    5: "Other"
}

# Add the simplified purpose variable
purpose_enc_col = encoded_vars.get('purpose', 'purpose_enc')
df_clean['purpose_simple'] = df_clean[purpose_enc_col].map(purpose_mapping)

# Check the new variable
print("\nPurpose simple distribution:")
purpose_counts = df_clean['purpose_simple'].value_counts().sort_index()
purpose_distribution = []
for code, count in purpose_counts.items():
    purpose_line = f"{code} ({purpose_labels.get(code, 'Unknown')}): {count}"
    purpose_distribution.append(purpose_line)
    print(purpose_line)

# Create a Word document for output
doc = Document()
doc.add_heading('Multilevel Truncated Negative Binomial Regression for Length of Stay Analysis', 0)
doc.add_heading('Data Preparation and Cleaning', level=1)

# Add missing data information
doc.add_paragraph('Missing Data Summary:')
missing_table = doc.add_table(rows=len(missing_data_summary)+1, cols=2)
missing_table.style = 'Table Grid'
missing_table.cell(0, 0).text = 'Variable'
missing_table.cell(0, 1).text = 'Missing Count'
for i, (var, count) in enumerate(missing_data_summary.items(), 1):
    missing_table.cell(i, 0).text = str(var)
    missing_table.cell(i, 1).text = str(count)

doc.add_paragraph('\nMissing Data Patterns:')
patterns_table = doc.add_table(rows=len(missing_patterns_counts)+1, cols=2)
patterns_table.style = 'Table Grid'
patterns_table.cell(0, 0).text = 'Number of Missing Variables'
patterns_table.cell(0, 1).text = 'Count'
for i, (pattern, count) in enumerate(missing_patterns_counts.items(), 1):
    patterns_table.cell(i, 0).text = str(pattern)
    patterns_table.cell(i, 1).text = str(count)

# Add Length of Stay histogram
doc.add_paragraph('\n')
doc.add_heading('Length of Stay Distribution', level=2)
doc.add_picture(los_hist_img, width=Inches(6))
doc.add_paragraph('Figure 1: Histogram of Length of Stay (Before Capping)')

# Add Capped LOS histogram
doc.add_paragraph('\n')
doc.add_heading('Capped Length of Stay Distribution', level=2)
doc.add_picture(los_capped_img, width=Inches(6))
doc.add_paragraph('Figure 2: Histogram of Length of Stay (After Capping at 95th Percentile)')

# Add LOS summary statistics
doc.add_paragraph('\n')
doc.add_heading('Length of Stay Summary Statistics', level=2)
los_stats_table = doc.add_table(rows=len(los_describe)+1, cols=2)
los_stats_table.style = 'Table Grid'
los_stats_table.cell(0, 0).text = 'Statistic'
los_stats_table.cell(0, 1).text = 'Value'
for i, (stat, value) in enumerate(los_describe.items(), 1):
    los_stats_table.cell(i, 0).text = str(stat)
    los_stats_table.cell(i, 1).text = f"{value:.4f}" if isinstance(value, (int, float)) else str(value)

# Add Purpose distribution
doc.add_paragraph('\n')
doc.add_heading('Purpose of Visit Distribution', level=2)
purpose_table = doc.add_table(rows=len(purpose_distribution)+1, cols=1)
purpose_table.style = 'Table Grid'
purpose_table.cell(0, 0).text = 'Purpose Category'
for i, purpose_text in enumerate(purpose_distribution, 1):
    purpose_table.cell(i, 0).text = purpose_text

# Fit simple negative binomial regression with continuous variables correctly specified
print("\nFitting simple negative binomial regression model...")
doc.add_paragraph('\n')
doc.add_heading('Negative Binomial Regression Model', level=1)

# Define continuous variables and create proper formula
continuous_vars = ['immigrant_population', 'import_from_slu', 'age', 'distance_miles', 
                   'state_percapita_income', 'state_unemployment']

# Make sure all continuous variables are properly formatted as numeric
for var in continuous_vars:
    if var in df_clean.columns:
        df_clean[var] = pd.to_numeric(df_clean[var], errors='coerce')

# Create formula with continuous variables properly treated
formula_parts = []
for var in continuous_vars:
    if var in df_clean.columns:
        formula_parts.append(var)

# Add categorical variables with proper C() notation
categorical_model_vars = ['sex_enc', 'marital_status_enc', 'employment_status_enc', 
                         'purpose_simple', 'accomd_type_enc', 'month_travel']

for var in categorical_model_vars:
    if var in df_clean.columns:
        # Use the encoded variable name or the original if available
        var_to_use = var
        formula_parts.append(f"C({var_to_use})")

# Combine into final formula
formula = 'los_capped ~ ' + ' + '.join(formula_parts)
print(f"Model formula: {formula}")

dummy_names = {}
for var in categorical_model_vars:
    if var in df_clean.columns:
        values = df_clean[var].unique()
        for val in values:
            if var == 'purpose_simple':
                dummy_names[f"{var}[T.{val}]"] = f"{var}_{purpose_labels.get(val, 'Unknown')}"
            else:
                dummy_names[f"{var}[T.{val}]"] = f"{var}_{val}"

# Add formula to document
doc.add_paragraph(f"Model formula: {formula}")

# Fit negative binomial model
nb_model = smf.glm(formula=formula, 
                  data=df_clean, 
                  family=sm.families.NegativeBinomial(link=sm.families.links.log()))

try:
    nb_results = nb_model.fit()
    print("\nNegative Binomial Regression Results:")
    summary_text = str(nb_results.summary())
    print(summary_text)
    
    # Add model summary to document
    doc.add_paragraph('\nModel Summary:')
    for line in summary_text.split('\n'):
        doc.add_paragraph(line)
    
    # Convert coefficients to incident rate ratios (IRR)
    print("\nIncident Rate Ratios (IRR):")
    irr = np.exp(nb_results.params)
    irr_conf = np.exp(nb_results.conf_int())
    irr_df = pd.DataFrame({'IRR': irr, 'Lower CI': irr_conf[0], 'Upper CI': irr_conf[1], 
                          'P-value': nb_results.pvalues})
    
    irr_df.index = [dummy_names.get(idx, idx) for idx in irr_df.index]
    print(irr_df)
    
    # Add IRR table to document
    doc.add_paragraph('\n')
    doc.add_heading('Incident Rate Ratios (IRR)', level=2)
    irr_table = doc.add_table(rows=len(irr_df)+1, cols=5)
    irr_table.style = 'Table Grid'
    irr_table.cell(0, 0).text = 'Variable'
    irr_table.cell(0, 1).text = 'IRR'
    irr_table.cell(0, 2).text = 'Lower CI'
    irr_table.cell(0, 3).text = 'Upper CI'
    irr_table.cell(0, 4).text = 'P-value'
    
    for i, (var, row) in enumerate(irr_df.iterrows(), 1):
        irr_table.cell(i, 0).text = str(var)
        irr_table.cell(i, 1).text = f"{row['IRR']:.4f}"
        irr_table.cell(i, 2).text = f"{row['Lower CI']:.4f}"
        irr_table.cell(i, 3).text = f"{row['Upper CI']:.4f}"
        irr_table.cell(i, 4).text = f"{row['P-value']:.4f}"
    
    # Predictions and diagnostics
    df_clean['predicted'] = nb_results.predict()
    df_clean['residuals'] = df_clean['los_capped'] - df_clean['predicted']
    
    # Plot residuals
    plt.figure(figsize=(10, 6))
    plt.scatter(df_clean['predicted'], df_clean['residuals'], alpha=0.5)
    plt.axhline(y=0, color='r', linestyle='-')
    plt.xlabel('Predicted Values')
    plt.ylabel('Residuals')
    plt.title('Residual Plot')
    plt.tight_layout()
    residuals_img = BytesIO()
    plt.savefig(residuals_img, format='png')
    residuals_img.seek(0)
    plt.close()
    
    # Add residuals plot to document
    doc.add_paragraph('\n')
    doc.add_heading('Diagnostics', level=2)
    doc.add_picture(residuals_img, width=Inches(6))
    doc.add_paragraph('Figure 3: Residuals Plot')
    
    # Check for heterogeneity across states if us_state is in the data
    if 'us_state_enc' in df_clean.columns or 'us_state' in df_clean.columns:
        state_var = 'us_state_enc' if 'us_state_enc' in df_clean.columns else 'us_state'
        state_means = df_clean.groupby(state_var)['los_capped'].mean().sort_values()
        
        plt.figure(figsize=(12, 8))
        state_means.plot(kind='bar')
        plt.xlabel('State')
        plt.ylabel('Average Length of Stay')
        plt.title('Mean Length of Stay by State')
        plt.xticks(rotation=90)
        plt.tight_layout()
        los_by_state_img = BytesIO()
        plt.savefig(los_by_state_img, format='png')
        los_by_state_img.seek(0)
        plt.close()
        
        # Add state analysis to document
        doc.add_paragraph('\n')
        doc.add_heading('State Analysis', level=2)
        doc.add_picture(los_by_state_img, width=Inches(6))
        doc.add_paragraph('Figure 4: Mean Length of Stay by State')
    
    # Approximated Multilevel Model
    doc.add_paragraph('\n')
    doc.add_heading('Approximated Multilevel Model', level=1)
    doc.add_paragraph('Using MixedLM to approximate a multilevel model with random effects for states.')
    
    
    
    if 'us_state_enc' in df_clean.columns:
        # For demonstration, we'll use a linear mixed model as an approximation
        # Prepare model variables
        y = df_clean['los_capped']
        
        # Create X matrix for fixed effects
        X_vars = []
        for var in continuous_vars:
            if var in df_clean.columns:
                X_vars.append(var)
        
        model_data = df_clean.dropna(subset=[*X_vars, 'los_capped', 'us_state_enc'])
        y = model_data['los_capped']
        X = model_data[X_vars].copy()
        
        # Add categorical variables (one-hot encoded)
        for var in categorical_model_vars:
            if var in df_clean.columns and var != 'us_state_enc':  # Exclude the grouping variable
                dummies = pd.get_dummies(df_clean[var], prefix=var, drop_first=True)
                X = pd.concat([X, dummies], axis=1)
        
        # Add intercept
        X = sm.add_constant(X)
        
        # Define groups for random effects
        groups =model_data['us_state_enc']
        
        # Fit mixed effects model
        mixed_model = MixedLM(y, X, groups)
        try:
            mixed_results = mixed_model.fit()
            mixed_summary = str(mixed_results.summary())
            print("\nApproximated Multilevel Model Results:")
            print(mixed_summary)
            
            # Add to document
            doc.add_paragraph('Model Summary:')
            for line in mixed_summary.split('\n'):
                doc.add_paragraph(line)
            
            # Add variance components
            doc.add_paragraph('\nVariance Components:')
            vc_table = doc.add_table(rows=3, cols=2)
            vc_table.style = 'Table Grid'
            vc_table.cell(0, 0).text = 'Component'
            vc_table.cell(0, 1).text = 'Estimate'
            vc_table.cell(1, 0).text = 'State Random Effect Variance'
            vc_table.cell(1, 1).text = f"{mixed_results.cov_re.iloc[0, 0]:.4f}"
            vc_table.cell(2, 0).text = 'Residual Variance'
            vc_table.cell(2, 1).text = f"{mixed_results.scale:.4f}"
            
            # Calculate intraclass correlation coefficient (ICC)
            state_var = mixed_results.cov_re.iloc[0, 0]
            residual_var = mixed_results.scale
            icc = state_var / (state_var + residual_var)
            
            doc.add_paragraph(f'\nIntraclass Correlation Coefficient (ICC): {icc:.4f}')
            doc.add_paragraph('The ICC represents the proportion of the total variance in length of stay ' +
                             'that is attributable to differences between states.')
            
        except Exception as e:
            error_msg = f"Error fitting mixed model: {str(e)}"
            print(error_msg)
            doc.add_paragraph(error_msg)
            doc.add_paragraph("The mixed effects model failed to converge. This can happen due to " +
                             "insufficient variation in the grouping variable or other model specification issues.")
    else:
        no_state_msg = "State variable not found for multilevel modeling."
        print(no_state_msg)
        doc.add_paragraph(no_state_msg)
    
except Exception as e:
    error_msg = f"\nError in model fitting: {str(e)}"
    print(error_msg)
    doc.add_paragraph(error_msg)
    doc.add_paragraph("You may need to check your data or consider using a different modeling approach.")

# Save the Word document
print("\nPlease select where to save the Word document...")
doc_path = select_file(
    "Save Analysis Report As", 
    [("Word Document", "*.docx"), ("All files", "*.*")],
    save=True
)

if doc_path:
    if not doc_path.endswith('.docx'):
        doc_path += '.docx'
    doc.save(doc_path)
    print(f"Analysis report saved to: {doc_path}")
else:
    print("Document not saved as no location was selected.")

print("\nAnalysis complete.")